"""Parse Word (.docx) email templates and extract {{variables}}.

Provides:
- Variable extraction from .docx files
- DOCX-to-HTML conversion preserving basic formatting
- Variable classification (simple vs. special/reserved)
- Column name normalization utility
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

# Variables that are generated programmatically, not expected as Excel columns
SPECIAL_VARIABLES: set[str] = {
    "module_table",
    "current_date",
}

# Pattern to match {{variable_name}} in template text.
# Allows spaces in variable names (e.g. {{First Name}}, {{Student ID}}).
VARIABLE_PATTERN = re.compile(r"\{\{([\w\s]+)\}\}")


def normalize_variable_name(name: str) -> str:
    """Normalize a column name to match template variable conventions.

    Rules: lowercase, strip whitespace, replace spaces with underscores.
    This ensures Excel column 'First Name' matches template {{first_name}}.
    """
    return str(name).strip().lower().replace(" ", "_")


@dataclass
class TableMapping:
    """Column definitions extracted from a DOCX table header row.

    Attributes:
        columns: Normalised column names extracted from header cells.
        has_data_table: True if the DOCX contains at least one table with a
                        detectable header row (bold first row).
    """

    columns: list[str] = field(default_factory=list)
    has_data_table: bool = False


@dataclass
class TemplateInfo:
    """Result of parsing a .docx template file.

    Attributes:
        html_content: The document converted to HTML with {{variables}} preserved.
        simple_variables: Variable names that map to Excel columns.
        special_variables: Variable names that are programmatically generated.
        table_mapping: If the DOCX contains a data table, the column definitions
                       extracted from its header row.
        all_variables: All variable names found (simple + special).
    """

    html_content: str
    simple_variables: list[str] = field(default_factory=list)
    special_variables: list[str] = field(default_factory=list)
    table_mapping: TableMapping | None = None

    @property
    def all_variables(self) -> list[str]:
        return self.simple_variables + self.special_variables


def _extract_variables_from_text(text: str) -> list[str]:
    """Return all unique {{variable}} names found in text, stripped of whitespace."""
    raw = VARIABLE_PATTERN.findall(text)
    return list(dict.fromkeys(v.strip() for v in raw))


def _classify_variables(variable_names: list[str]) -> tuple[list[str], list[str]]:
    """Split variable names into (simple, special) lists.

    Simple variables map to Excel columns.
    Special variables are generated programmatically (e.g., module_table).
    Variable names are normalised before classification so that spaced
    variants like ``{{Module Table}}`` still match the reserved set.
    """
    simple = []
    special = []
    for name in variable_names:
        normalized = normalize_variable_name(name)
        if normalized in SPECIAL_VARIABLES:
            special.append(normalized)
        else:
            simple.append(normalized)
    return simple, special


def _run_to_html(run) -> str:
    """Convert a single python-docx Run to inline HTML.

    Preserves bold (<strong>) and italic (<em>) formatting.
    """
    text = run.text
    if not text:
        return ""

    if run.bold and run.italic:
        return f"<strong><em>{text}</em></strong>"
    elif run.bold:
        return f"<strong>{text}</strong>"
    elif run.italic:
        return f"<em>{text}</em>"
    return text


def _normalize_placeholders(html: str) -> str:
    """Replace {{Variable Name}} with {{variable_name}} in HTML so Jinja2 can render.

    Jinja2 identifiers cannot contain spaces.  This rewrites every {{…}}
    placeholder by normalising the inner variable name (lowercase, spaces→underscores).
    """
    def _replace(match: re.Match[str]) -> str:
        inner = match.group(1).strip()
        return "{{" + normalize_variable_name(inner) + "}}"
    return VARIABLE_PATTERN.sub(_replace, html)


def _is_header_row(row) -> bool:
    """Return True if *all* non-empty cells in the row are bold.

    A header row is detected when every cell that contains text has at least
    one bold run.  Empty cells are ignored for this check.
    """
    has_text = False
    for cell in row.cells:
        cell_text = cell.text.strip()
        if not cell_text:
            continue
        has_text = True
        # Check if any paragraph in the cell has a bold run
        any_bold = any(
            run.bold
            for para in cell.paragraphs
            for run in para.runs
            if run.text.strip()
        )
        if not any_bold:
            return False
    return has_text  # at least one non-empty cell AND all non-empty cells bold


def _extract_header_columns(table) -> list[str]:
    """Return normalised column names from the first row of a table.

    Only called when ``_is_header_row(table.rows[0])`` is True.
    """
    columns: list[str] = []
    for cell in table.rows[0].cells:
        raw = cell.text.strip()
        if raw:
            columns.append(normalize_variable_name(raw))
        else:
            columns.append("")  # preserve column position for empty headers
    return columns


def _build_data_table_html(table, header_columns: list[str]) -> tuple[str, list[str]]:
    """Convert a DOCX table to HTML with Jinja2 ``{% for %}`` data rows.

    The header row is kept as static HTML.  All subsequent rows are replaced
    with a single Jinja2 loop that iterates ``modules`` and calls
    ``module.get(column_name, '')`` for each header column.

    Returns:
        (html_string, list_of_variables_found)
    """
    vars_found: list[str] = []
    rows_html: list[str] = []

    # Header row — keep as static HTML
    header_cells: list[str] = []
    for cell in table.rows[0].cells:
        parts: list[str] = []
        for para in cell.paragraphs:
            for run in para.runs:
                if run.text:
                    vars_found.extend(_extract_variables_from_text(run.text))
                    parts.append(_run_to_html(run))
        header_cells.append("<th>" + "".join(parts) + "</th>")
    rows_html.append("<tr>" + "".join(header_cells) + "</tr>")

    # Data rows — Jinja2 for loop
    data_cells: list[str] = []
    for col_name in header_columns:
        if col_name:
            data_cells.append(
                "<td>{{ module.get('" + col_name + "', '') }}</td>"
            )
        else:
            data_cells.append("<td></td>")

    rows_html.append(
        "{% for module in modules %}<tr>"
        + "".join(data_cells)
        + "</tr>{% endfor %}"
    )

    html = '<table border="1" cellpadding="6">' + "".join(rows_html) + "</table>"
    return (html, vars_found)


def parse_and_convert(docx_path: Path) -> TemplateInfo:
    """Parse a .docx file, extract variables, and convert content to HTML.

    Args:
        docx_path: Path to the .docx template file.

    Returns:
        TemplateInfo with html_content, simple_variables, and special_variables.

    Paragraphs become <p> tags. Tables become <table> with <tr>/<td>.
    Bold runs wrap in <strong>, italic runs wrap in <em>.
    {{variables}} are preserved as literal text for Jinja2 rendering.
    """
    # Lazy import so the module can be imported without python-docx installed
    from docx import Document

    doc = Document(str(docx_path))
    all_vars: list[str] = []
    html_parts: list[str] = []
    table_mapping: TableMapping | None = None

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is not None:
                para_html, para_vars = _paragraph_to_html(para)
                if para_html.strip():
                    html_parts.append(para_html)
                    all_vars.extend(para_vars)

        elif tag == "tbl":
            table = _find_table(doc, element)
            if table is None:
                continue

            # Detect data tables: first row bold → header → dynamic rows
            if len(table.rows) >= 2 and _is_header_row(table.rows[0]):
                header_columns = _extract_header_columns(table)
                table_html, table_vars = _build_data_table_html(table, header_columns)
                if table_mapping is None:
                    table_mapping = TableMapping(
                        columns=header_columns,
                        has_data_table=True,
                    )
            else:
                table_html, table_vars = _table_to_html(table)

            if table_html.strip():
                html_parts.append(table_html)
                all_vars.extend(table_vars)

    if not html_parts:
        html_parts.append("<p></p>")

    raw_html = "\n".join(html_parts)
    # Normalize {{variable names with spaces}} → {{variable_names_with_underscores}}
    # so Jinja2 can render them (Jinja2 identifiers cannot contain spaces).
    html_content = _normalize_placeholders(raw_html)

    simple, special = _classify_variables(list(dict.fromkeys(all_vars)))
    return TemplateInfo(
        html_content=html_content,
        simple_variables=simple,
        special_variables=special,
        table_mapping=table_mapping,
    )


def _find_paragraph(doc, element):
    """Find the Paragraph object matching an XML element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    """Find the Table object matching an XML element."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _paragraph_to_html(para) -> tuple[str, list[str]]:
    """Convert a Paragraph to HTML <p> tag with inline formatting.

    Returns:
        (html_string, list_of_variables_found)
    """
    vars_found: list[str] = []
    run_parts: list[str] = []

    for run in para.runs:
        run_text = run.text
        if run_text:
            vars_found.extend(_extract_variables_from_text(run_text))
            run_parts.append(_run_to_html(run))

    if not run_parts:
        return ("", [])

    html = "<p>" + "".join(run_parts) + "</p>"
    return (html, vars_found)


def _table_to_html(table) -> tuple[str, list[str]]:
    """Convert a Table to an HTML <table> with <tr> and <td>.

    Returns:
        (html_string, list_of_variables_found)
    """
    vars_found: list[str] = []
    rows_html: list[str] = []

    for row in table.rows:
        cells_html: list[str] = []
        for cell in row.cells:
            cell_text_parts: list[str] = []
            for para in cell.paragraphs:
                for run in para.runs:
                    run_text = run.text
                    if run_text:
                        vars_found.extend(_extract_variables_from_text(run_text))
                        cell_text_parts.append(_run_to_html(run))
            cells_html.append("<td>" + "".join(cell_text_parts) + "</td>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")

    if not rows_html:
        return ("", [])

    html = '<table border="1" cellpadding="6">' + "".join(rows_html) + "</table>"
    return (html, vars_found)


def extract_variables(docx_path: Path) -> TemplateInfo:
    """Extract variables from a .docx file without full HTML conversion.

    Faster than parse_and_convert; used when only variable names are needed.
    Also detects data tables and extracts column mappings.

    Args:
        docx_path: Path to the .docx template file.

    Returns:
        TemplateInfo with empty html_content and extracted variable lists.
    """
    from docx import Document

    doc = Document(str(docx_path))
    all_vars: list[str] = []
    table_mapping: TableMapping | None = None

    for para in doc.paragraphs:
        all_vars.extend(_extract_variables_from_text(para.text))

    for table in doc.tables:
        # Detect data table header
        if len(table.rows) >= 2 and _is_header_row(table.rows[0]):
            header_columns = _extract_header_columns(table)
            if table_mapping is None:
                table_mapping = TableMapping(
                    columns=header_columns,
                    has_data_table=True,
                )
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_vars.extend(_extract_variables_from_text(para.text))

    simple, special = _classify_variables(list(dict.fromkeys(all_vars)))
    return TemplateInfo(
        html_content="",
        simple_variables=simple,
        special_variables=special,
        table_mapping=table_mapping,
    )
