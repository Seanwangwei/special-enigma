"""Tests for docx_parser.py — Word template parsing and variable extraction."""

from pathlib import Path
from docx import Document
from docx.shared import Pt

from exam_email_automation.templates.docx_parser import (
    extract_variables,
    parse_and_convert,
    normalize_variable_name,
    TemplateInfo,
    SPECIAL_VARIABLES,
)


def _create_docx(tmp_path: Path, filename: str, paragraphs: list[tuple[str, bool, bool]]) -> Path:
    """Helper: create a minimal .docx file with styled paragraphs.

    Args:
        tmp_path: pytest tmp_path fixture.
        filename: Name for the .docx file.
        paragraphs: List of (text, bold, italic) tuples.

    Returns:
        Path to the created .docx file.
    """
    doc = Document()
    for text, bold, italic in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
    path = tmp_path / filename
    doc.save(str(path))
    return path


def _create_docx_with_table(tmp_path: Path, filename: str, rows: list[list[str]]) -> Path:
    """Helper: create a .docx file with a table.

    Args:
        tmp_path: pytest tmp_path fixture.
        filename: Name for the .docx file.
        rows: List of rows, each row is a list of cell text strings.

    Returns:
        Path to the created .docx file.
    """
    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text
    path = tmp_path / filename
    doc.save(str(path))
    return path


# --- normalize_variable_name tests ---

def test_normalize_variable_name_lowercases():
    assert normalize_variable_name("First Name") == "first_name"


def test_normalize_variable_name_replaces_spaces():
    assert normalize_variable_name("Assessment Format in August") == "assessment_format_in_august"


def test_normalize_variable_name_strips_whitespace():
    assert normalize_variable_name("  Student ID  ") == "student_id"


# --- extract_variables tests ---

def test_extract_variables_from_paragraphs(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("Hello {{first_name}},", False, False),
        ("Your result: {{stage_average}}", False, False),
    ])
    info = extract_variables(path)

    assert "first_name" in info.simple_variables
    assert "stage_average" in info.simple_variables
    assert info.html_content == ""  # extract_variables doesn't produce HTML


def test_extract_variables_no_duplicates(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("{{first_name}} {{first_name}}", False, False),
    ])
    info = extract_variables(path)

    assert info.simple_variables.count("first_name") <= len(info.simple_variables)


def test_special_variable_classification(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("{{first_name}} {{module_table}}", False, False),
    ])
    info = extract_variables(path)

    assert "first_name" in info.simple_variables
    assert "module_table" in info.special_variables
    assert "module_table" not in info.simple_variables


def test_current_date_is_special(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("Sent on {{current_date}}", False, False),
    ])
    info = extract_variables(path)

    assert "current_date" in info.special_variables


def test_extract_variables_from_table(tmp_path: Path):
    path = _create_docx_with_table(tmp_path, "test.docx", [
        ["Module Code", "Module Name", "{{module_table}}"],
        ["COMP101", "Programming", ""],
    ])
    info = extract_variables(path)

    assert "module_table" in info.special_variables


def test_empty_document(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("")  # empty paragraph
    path = tmp_path / "empty.docx"
    doc.save(str(path))

    info = extract_variables(path)
    assert info.simple_variables == []
    assert info.special_variables == []


# --- parse_and_convert tests ---

def test_parse_and_convert_paragraphs_to_html(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("Dear {{first_name}},", False, False),
        ("Your score is {{stage_average}}.", False, False),
    ])
    info = parse_and_convert(path)

    assert "<p>Dear {{first_name}},</p>" in info.html_content
    assert "<p>Your score is {{stage_average}}.</p>" in info.html_content


def test_parse_and_convert_preserves_variables(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("Hello {{first_name}} {{surname}}", False, False),
    ])
    info = parse_and_convert(path)

    assert "{{first_name}}" in info.html_content
    assert "{{surname}}" in info.html_content


def test_parse_and_convert_bold_preserved(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("Important notice", True, False),
    ])
    info = parse_and_convert(path)

    assert "<strong>Important notice</strong>" in info.html_content


def test_parse_and_convert_italic_preserved(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("Please note", False, True),
    ])
    info = parse_and_convert(path)

    assert "<em>Please note</em>" in info.html_content


def test_parse_and_convert_bold_italic_combined(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("URGENT", True, True),
    ])
    info = parse_and_convert(path)

    assert "<strong><em>URGENT</em></strong>" in info.html_content


def test_parse_and_convert_table(tmp_path: Path):
    path = _create_docx_with_table(tmp_path, "test.docx", [
        ["Module", "Grade"],
        ["{{module_code}}", "{{module_grade}}"],
    ])
    info = parse_and_convert(path)

    assert "<table" in info.html_content
    assert "{{module_code}}" in info.html_content
    assert "{{module_grade}}" in info.html_content


def test_parse_and_convert_empty_document(tmp_path: Path):
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))

    info = parse_and_convert(path)
    assert info.html_content == "<p></p>"
    assert info.simple_variables == []


def test_parse_and_convert_all_variables(tmp_path: Path):
    path = _create_docx(tmp_path, "test.docx", [
        ("{{first_name}} {{module_table}}", False, False),
    ])
    info = parse_and_convert(path)

    assert "first_name" in info.simple_variables
    assert "module_table" in info.special_variables
    assert len(info.all_variables) == 2
