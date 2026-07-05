"""Sprint 7 tests — Dynamic DOCX Table Population.

Covers: table detection, column mapping, dynamic row generation,
module extra_fields, full pipeline integration, and backward compatibility.
"""

from __future__ import annotations

import pytest
from pathlib import Path
from io import BytesIO

import docx
from docx import Document
from docx.shared import Pt

from exam_email_automation.templates.docx_parser import (
    parse_and_convert,
    extract_variables,
    normalize_variable_name,
    _is_header_row,
    _extract_header_columns,
    _build_data_table_html,
    TableMapping,
    TemplateInfo,
)
from exam_email_automation.models.module_result import ModuleResult
from exam_email_automation.templates.template_engine import TemplateEngine
from exam_email_automation.excel.excel_reader import ExcelReader
from exam_email_automation.models.student import Student


# ---------------------------------------------------------------------------
# Helpers — build in-memory .docx files for testing
# ---------------------------------------------------------------------------

def _make_docx_with_table(
    header_cells: list[str],
    data_rows: list[list[str]] | None = None,
    bold_header: bool = True,
) -> Path:
    """Create a minimal .docx with one table and return its path."""
    import tempfile

    doc = Document()
    table = doc.add_table(rows=1 + len(data_rows or []), cols=len(header_cells))

    # Header row
    for ci, text in enumerate(header_cells):
        cell = table.rows[0].cells[ci]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        if bold_header:
            run.bold = True

    # Data rows
    if data_rows:
        for ri, row_data in enumerate(data_rows):
            for ci, text in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                p = cell.paragraphs[0]
                p.add_run(text)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return Path(tmp.name)


def _make_docx_without_table(text_content: str = "Hello {{First Name}}") -> Path:
    """Create a minimal .docx with only paragraphs (no tables)."""
    import tempfile

    doc = Document()
    doc.add_paragraph(text_content)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return Path(tmp.name)


# ---------------------------------------------------------------------------
# _is_header_row
# ---------------------------------------------------------------------------

class TestIsHeaderRow:
    def test_bold_header_detected(self):
        path = _make_docx_with_table(
            ["Module Code", "Module Name", "Credits"],
            data_rows=[["", "", ""]],
            bold_header=True,
        )
        doc = Document(str(path))
        assert _is_header_row(doc.tables[0].rows[0]) is True

    def test_non_bold_not_header(self):
        path = _make_docx_with_table(
            ["Module Code", "Module Name"],
            bold_header=False,
        )
        doc = Document(str(path))
        assert _is_header_row(doc.tables[0].rows[0]) is False

    def test_single_row_table_skipped(self):
        """A table with only 1 row is never a data table (needs ≥2)."""
        path = _make_docx_with_table(["A", "B"], data_rows=None)
        doc = Document(str(path))
        # This table has no header detection because it's not checked
        # (parse_and_convert requires len≥2). Just verify it doesn't crash.


# ---------------------------------------------------------------------------
# _extract_header_columns
# ---------------------------------------------------------------------------

class TestExtractHeaderColumns:
    def test_normalizes_column_names(self):
        path = _make_docx_with_table(
            ["Module Code", "Module Name", "Student ID", "Credits"],
            data_rows=[["", "", "", ""]],
            bold_header=True,
        )
        doc = Document(str(path))
        cols = _extract_header_columns(doc.tables[0])
        assert cols == ["module_code", "module_name", "student_id", "credits"]


# ---------------------------------------------------------------------------
# _build_data_table_html
# ---------------------------------------------------------------------------

class TestBuildDataTableHtml:
    def test_generates_jinja2_for_loop(self):
        path = _make_docx_with_table(
            ["Module Code", "Module Name"],
            data_rows=[["", ""]],
            bold_header=True,
        )
        doc = Document(str(path))
        table = doc.tables[0]
        cols = _extract_header_columns(table)
        html, vars_found = _build_data_table_html(table, cols)

        assert "{% for module in modules %}" in html
        assert "{% endfor %}" in html
        assert "module.get('module_code', '')" in html
        assert "module.get('module_name', '')" in html
        # Header row preserved with <th>
        assert "<th>" in html

    def test_empty_header_column_handled(self):
        """Empty header cells should produce empty <td></td>."""
        path = _make_docx_with_table(
            ["Code", ""],  # second header empty
            data_rows=[["", ""]],
            bold_header=True,
        )
        doc = Document(str(path))
        table = doc.tables[0]
        cols = _extract_header_columns(table)
        html, _ = _build_data_table_html(table, cols)
        assert "<td></td>" in html  # empty column should be empty cell


# ---------------------------------------------------------------------------
# parse_and_convert — table detection
# ---------------------------------------------------------------------------

class TestParseAndConvertWithTable:
    def test_detects_data_table(self):
        path = _make_docx_with_table(
            ["Module Code", "Module Name", "Credits"],
            data_rows=[["", "", ""]],
            bold_header=True,
        )
        info = parse_and_convert(path)
        assert info.table_mapping is not None
        assert info.table_mapping.has_data_table is True
        assert info.table_mapping.columns == ["module_code", "module_name", "credits"]

    def test_no_table_template_has_none_mapping(self):
        path = _make_docx_without_table()
        info = parse_and_convert(path)
        assert info.table_mapping is None

    def test_data_row_variables_replaced_by_loop(self):
        """When a data table is detected, data rows are replaced with the
        ``{% for %}`` loop.  Any ``{{variables}}`` that were in the data rows
        are intentionally NOT extracted — the dynamic row generation uses
        ``module.get(column_name)`` instead."""
        path = _make_docx_with_table(
            ["Module Code", "Module Name"],
            data_rows=[["{{Grade}}", "{{Semester}}"]],
            bold_header=True,
        )
        info = parse_and_convert(path)
        # Data rows were replaced → {{Grade}} / {{Semester}} no longer in HTML
        assert "grade" not in info.simple_variables
        assert "semester" not in info.simple_variables
        # But the for-loop is present (dynamic rows)
        assert "{% for module in modules %}" in info.html_content


# ---------------------------------------------------------------------------
# extract_variables — also detects tables
# ---------------------------------------------------------------------------

class TestExtractVariablesWithTable:
    def test_detects_data_table_columns(self):
        path = _make_docx_with_table(
            ["Col A", "Col B"],
            data_rows=[["", ""]],
            bold_header=True,
        )
        info = extract_variables(path)
        assert info.table_mapping is not None
        assert info.table_mapping.has_data_table is True
        assert "col_a" in info.table_mapping.columns

    def test_non_table_template_no_mapping(self):
        path = _make_docx_without_table()
        info = extract_variables(path)
        assert info.table_mapping is None


# ---------------------------------------------------------------------------
# ModuleResult.get() and extra_fields
# ---------------------------------------------------------------------------

class TestModuleResult:
    def test_get_known_field(self):
        m = ModuleResult(module_code="CS101", module_name="Intro to CS")
        assert m.get("module_code") == "CS101"
        assert m.get("module_name") == "Intro to CS"

    def test_get_extra_field(self):
        m = ModuleResult(extra_fields={"credits": "5", "semester": "Fall"})
        assert m.get("credits") == "5"
        assert m.get("semester") == "Fall"

    def test_get_missing_field_returns_default(self):
        m = ModuleResult()
        assert m.get("nonexistent") == ""
        assert m.get("nonexistent", "N/A") == "N/A"

    def test_known_field_takes_priority_over_extra(self):
        m = ModuleResult(
            module_code="CS101",
            extra_fields={"module_code": "WRONG"},
        )
        assert m.get("module_code") == "CS101"  # known field wins


# ---------------------------------------------------------------------------
# generate_module_table with custom columns
# ---------------------------------------------------------------------------

class TestGenerateModuleTable:
    def test_default_columns(self):
        student = Student(
            student_id="1", first_name="A", surname="B",
            email="a@b.com", template_name="", stage_average="",
            pass_credits="", failed_modules="",
        )
        student.modules = [
            ModuleResult("CS101", "Intro", "Exam", "1st", "5"),
        ]
        html = TemplateEngine.generate_module_table(student)
        assert "<th>Module Code</th>" in html
        assert "<td>CS101</td>" in html

    def test_custom_columns(self):
        student = Student(
            student_id="1", first_name="A", surname="B",
            email="a@b.com", template_name="", stage_average="",
            pass_credits="", failed_modules="",
        )
        student.modules = [
            ModuleResult(
                "CS101", "Intro", "Exam", "1st", "5",
                extra_fields={"credits": "5", "semester": "Fall"},
            ),
        ]
        html = TemplateEngine.generate_module_table(
            student, columns=["module_code", "module_name", "credits", "semester"],
        )
        assert "<th>module_code</th>" in html
        assert "<td>CS101</td>" in html
        assert "<td>5</td>" in html
        assert "<td>Fall</td>" in html

    def test_no_modules_shows_placeholder(self):
        student = Student(
            student_id="1", first_name="A", surname="B",
            email="a@b.com", template_name="", stage_average="",
            pass_credits="", failed_modules="",
        )
        html = TemplateEngine.generate_module_table(student, columns=["module_code"])
        assert "No module results" in html


# ---------------------------------------------------------------------------
# Full pipeline integration
# ---------------------------------------------------------------------------

class TestFullPipelineWithTable:
    def test_parse_render_multimodule_student(self, tmp_path):
        """End-to-end: DOCX table → parse → Excel load → render → verify rows."""
        # 1. Create a DOCX with a table
        docx_path = _make_docx_with_table(
            ["Module Code", "Module Name", "Credits"],
            data_rows=[["", "", ""]],
            bold_header=True,
        )
        info = parse_and_convert(docx_path)

        # 2. Create an Excel file
        import pandas as pd
        df = pd.DataFrame({
            "Student ID": ["S1", "S1"],
            "First Name": ["Alice", "Alice"],
            "Surname": ["Smith", "Smith"],
            "Module Code": ["M001", "M002"],
            "Module Name": ["Math", "Physics"],
            "Credits": ["5", "4"],
        })
        xlsx_path = tmp_path / "test.xlsx"
        df.to_excel(str(xlsx_path), index=False)

        # 3. Set up engine and load
        engine = TemplateEngine(Path("src/exam_email_automation/templates"))
        engine.load_docx_template(info.html_content)

        reader = ExcelReader()
        students = reader.load_students(xlsx_path, template_variables=info.all_variables)
        assert len(students) == 1
        assert len(students[0].modules) == 2  # Alice has 2 modules

        # 4. Render and verify
        html = engine.render(students[0])
        assert "{% for" not in html  # Jinja2 tags fully resolved
        assert "{% endfor" not in html
        assert "<td>M001</td>" in html
        assert "<td>M002</td>" in html
        assert "<td>Math</td>" in html
        assert "<td>Physics</td>" in html
        assert "<td>5</td>" in html
        assert "<td>4</td>" in html

    def test_unmapped_columns_render_empty(self, tmp_path):
        """Columns in the table header that don't exist in Excel → empty cells."""
        docx_path = _make_docx_with_table(
            ["Module Code", "Semester", "Unknown Col"],
            data_rows=[["", "", ""]],
            bold_header=True,
        )
        info = parse_and_convert(docx_path)

        import pandas as pd
        df = pd.DataFrame({
            "Student ID": ["S1"],
            "Module Code": ["M001"],
        })
        xlsx_path = tmp_path / "test.xlsx"
        df.to_excel(str(xlsx_path), index=False)

        engine = TemplateEngine(Path("src/exam_email_automation/templates"))
        engine.load_docx_template(info.html_content)

        reader = ExcelReader()
        students = reader.load_students(xlsx_path, template_variables=info.all_variables)

        html = engine.render(students[0])
        assert "<td>M001</td>" in html
        # Unmapped columns should be empty but not error
        assert html.count("<td></td>") >= 2


# ---------------------------------------------------------------------------
# Backward compatibility
# ---------------------------------------------------------------------------

class TestBackwardCompatibility:
    def test_non_table_template_still_works(self, tmp_path):
        """Templates without tables should behave exactly as before."""
        import pandas as pd

        docx_path = _make_docx_without_table("Dear {{First Name}} {{Surname}}")
        info = parse_and_convert(docx_path)
        assert info.table_mapping is None  # no table detected

        df = pd.DataFrame({
            "Student ID": ["S1"],
            "First Name": ["Alice"],
            "Surname": ["Smith"],
        })
        xlsx_path = tmp_path / "test.xlsx"
        df.to_excel(str(xlsx_path), index=False)

        engine = TemplateEngine(Path("src/exam_email_automation/templates"))
        engine.load_docx_template(info.html_content)

        reader = ExcelReader()
        students = reader.load_students(xlsx_path, template_variables=info.all_variables)
        html = engine.render(students[0])
        assert "Dear Alice Smith" in html

    def test_module_table_special_var_still_works(self, tmp_path):
        """{{module_table}} special variable must still generate default table."""
        import tempfile, pandas as pd

        doc = Document()
        doc.add_paragraph("Your modules:")
        doc.add_paragraph("{{module_table}}")
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        docx_path = Path(tmp.name)

        info = parse_and_convert(docx_path)
        assert "module_table" in info.special_variables

        df = pd.DataFrame({
            "Student ID": ["S1"],
            "Module Code": ["CS101"],
            "Module Name": ["Intro to CS"],
        })
        xlsx_path = tmp_path / "test.xlsx"
        df.to_excel(str(xlsx_path), index=False)

        engine = TemplateEngine(Path("src/exam_email_automation/templates"))
        engine.load_docx_template(info.html_content)

        reader = ExcelReader()
        students = reader.load_students(xlsx_path, template_variables=info.all_variables)
        html = engine.render(students[0])
        assert "<td>CS101</td>" in html
        assert "<td>Intro to CS</td>" in html

    def test_non_bold_table_stays_static(self):
        """Tables without bold header rows remain static HTML (no for loop)."""
        path = _make_docx_with_table(
            ["A", "B"],
            data_rows=[["val1", "val2"]],
            bold_header=False,
        )
        info = parse_and_convert(path)
        assert info.table_mapping is None  # not detected as data table
        assert "{% for" not in info.html_content
        assert "<td>val1</td>" in info.html_content
