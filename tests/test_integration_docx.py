"""Integration tests: .docx → parse → validate → render → send pipeline."""

from pathlib import Path
from docx import Document

from exam_email_automation.templates.docx_parser import parse_and_convert
from exam_email_automation.templates.template_engine import TemplateEngine
from exam_email_automation.templates.template_metadata import TemplateMeta, load_companion_yaml
from exam_email_automation.validation.validator import Validator, MismatchReport
from exam_email_automation.email.email_builder import EmailBuilder
from exam_email_automation.models.student import Student


def _create_docx_with_vars(tmp_path: Path, paragraphs: list[str]) -> Path:
    """Helper: create a minimal .docx with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    path = tmp_path / "template.docx"
    doc.save(str(path))
    return path


def _create_yaml(tmp_path: Path, subject: str) -> Path:
    """Helper: create a companion YAML file."""
    import yaml
    yaml_path = tmp_path / "template.yaml"
    with open(yaml_path, "w") as f:
        yaml.dump({"subject": subject}, f)
    return yaml_path


# --- Integration: parse → template engine → render ---

def test_full_pipeline_parse_and_render(tmp_path: Path):
    """A parsed .docx template should render correctly with student data."""
    docx_path = _create_docx_with_vars(tmp_path, [
        "Dear {{first_name}} {{surname}},",
        "Your stage average is {{stage_average}}.",
        "{{module_table}}",
    ])

    template_info = parse_and_convert(docx_path)
    assert "first_name" in template_info.simple_variables
    assert "surname" in template_info.simple_variables
    assert "stage_average" in template_info.simple_variables
    assert "module_table" in template_info.special_variables
    assert "<p>" in template_info.html_content  # converted to HTML
    assert "{{first_name}}" in template_info.html_content  # vars preserved

    # Load into TemplateEngine
    engine = TemplateEngine(tmp_path)  # empty folder for FileSystemLoader
    engine.load_docx_template(template_info.html_content)
    assert engine.has_uploaded_template()

    # Render with a student
    student = Student(
        student_id="S001", first_name="Alice", surname="Smith",
        email="alice@test.com", template_name="Template 1",
        stage_average="72", pass_credits="45", failed_modules="0",
        extra_fields={"programme_name": "Computer Science"},
    )
    html = engine.render(student)

    assert "Dear Alice Smith," in html
    assert "72" in html
    # Clean up
    engine.unload_uploaded_template()
    assert not engine.has_uploaded_template()


# --- Integration: validate_template_against_excel ---

def test_validate_template_against_matching_excel():
    """Template variables matching Excel columns should produce a clean report."""
    template_vars = ["first_name", "surname", "stage_average", "module_table"]
    excel_headers = ["Student ID", "First Name", "Surname", "Email", "Stage Average"]

    report = Validator.validate_template_against_excel(template_vars, excel_headers)

    assert report.is_clean
    assert "first_name" not in report.missing_in_excel
    assert "surname" not in report.missing_in_excel


def test_validate_template_against_mismatched_excel():
    """Missing template variables should be reported."""
    template_vars = ["first_name", "surname", "programme_name"]
    excel_headers = ["First Name", "Email"]  # missing surname and programme_name

    report = Validator.validate_template_against_excel(template_vars, excel_headers)

    assert not report.is_clean
    assert "surname" in report.missing_in_excel
    assert "programme_name" in report.missing_in_excel
    # "Email" should be in unused columns (not in template vars)
    assert any("Email" in c or "email" in c.lower() for c in report.unused_excel_columns)


def test_special_variables_excluded_from_validation():
    """Special variables like module_table should be excluded from Excel validation."""
    template_vars = ["first_name", "module_table"]
    excel_headers = ["First Name"]

    report = Validator.validate_template_against_excel(template_vars, excel_headers)

    assert report.is_clean  # module_table is special, not checked
    assert "module_table" not in report.missing_in_excel


# --- Integration: EmailBuilder with override subject ---

def test_email_builder_override_subject():
    """When override_subject is set, it should replace SUBJECT_MAP lookup."""
    builder = EmailBuilder(override_subject="Custom Subject Line")
    student = Student(
        student_id="S001", first_name="Alice", surname="Smith",
        email="alice@test.com", template_name="Template 1",
        stage_average="72", pass_credits="45", failed_modules="0",
    )
    msg = builder.compose(student, "<p>Test</p>")

    assert msg.subject == "Custom Subject Line"
    assert msg.to_address == "alice@test.com"
    assert msg.html_body == "<p>Test</p>"


def test_email_builder_fallback_to_subject_map():
    """Without override_subject, should fall back to SUBJECT_MAP."""
    builder = EmailBuilder(override_subject=None)
    student = Student(
        student_id="S001", first_name="Alice", surname="Smith",
        email="alice@test.com", template_name="Template 1",
        stage_average="72", pass_credits="45", failed_modules="0",
    )
    msg = builder.compose(student, "<p>Test</p>")

    assert msg.subject == "Exam Results and Progression"


# --- Integration: Student extra_fields → to_context() → template render ---

def test_extra_fields_flow_to_template(tmp_path: Path):
    """Extra fields from Excel should flow through to template rendering."""
    docx_path = _create_docx_with_vars(tmp_path, [
        "Programme: {{programme_name}}",
    ])
    template_info = parse_and_convert(docx_path)

    engine = TemplateEngine(tmp_path)
    engine.load_docx_template(template_info.html_content)

    student = Student(
        student_id="S002", first_name="Bob", surname="Jones",
        email="bob@test.com", template_name="Template 1",
        stage_average="55", pass_credits="30", failed_modules="1",
        extra_fields={"programme_name": "Physics"},
    )
    html = engine.render(student)

    assert "Programme: Physics" in html
    engine.unload_uploaded_template()


# --- Integration: companion YAML loading ---

def test_load_companion_yaml_found(tmp_path: Path):
    """A .yaml file next to the .docx should be found and parsed."""
    # Create both .docx and .yaml
    doc = Document()
    doc.add_paragraph("{{first_name}}")
    docx_path = tmp_path / "template.docx"
    doc.save(str(docx_path))

    import yaml
    yaml_path = tmp_path / "template.yaml"
    with open(yaml_path, "w") as f:
        yaml.dump({"subject": "Your Results"}, f)

    data = load_companion_yaml(docx_path)
    assert data is not None
    assert data["subject"] == "Your Results"


def test_load_companion_yaml_not_found(tmp_path: Path):
    """When no companion YAML exists, should return None."""
    doc = Document()
    doc.add_paragraph("{{first_name}}")
    docx_path = tmp_path / "standalone.docx"
    doc.save(str(docx_path))

    data = load_companion_yaml(docx_path)
    assert data is None


# --- MismatchReport properties ---

def test_mismatch_report_is_clean_when_empty():
    report = MismatchReport(missing_in_excel=[], unused_excel_columns=[])
    assert report.is_clean


def test_mismatch_report_not_clean_with_missing():
    report = MismatchReport(
        missing_in_excel=["first_name"],
        unused_excel_columns=[],
    )
    assert not report.is_clean
